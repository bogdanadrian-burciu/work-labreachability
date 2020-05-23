# Checking ICMP reachability to devices referenced by Excel file in format (Node_IP_address, Node_name, Node_owner)
# Task automation: check if node name contains "VM" and if not ping it, if not reachable write a row in a Word file and color that with a unique color per node responsible
# Bogdan Adrian Burciu 23/05/2020 vers 100

# -------------------------
# Credits:
# https://dev.to/ankitdobhal/let-s-ping-the-network-with-python-scapy-5g18
# https://stackoverflow.com/questions/25980777/new-to-scapy-trying-to-understand-the-sr
# https://thepacketgeek.com/scapy-p-06-sending-and-receiving-with-scapy/
# https://stackoverflow.com/questions/14931906/extract-columns-from-excel-using-python
# https://stackoverflow.com/questions/3389574/check-if-multiple-strings-exist-in-another-string
# https://stackoverflow.com/questions/41979095/write-text-in-particular-font-color-in-ms-word-using-python-docx

import docx
import xlrd
from scapy.all import *
from docx.shared import RGBColor

labxlsxAbsPath =  input("Please input absolute path to the Excel file containing lab devices in format (Node_IP_address, Node_name, Node_owner), be careful to have only valid IPs (like C:\\Users\\boburciu\\Desktop\\Tier2 Procedures\\4thFloorLab.xlsx ):")
sh = xlrd.open_workbook(labxlsxAbsPath).sheet_by_index(0)
ips = sh.col_values(0, start_rowx=0)
names = sh.col_values(1, start_rowx=0)
reps = sh.col_values(2, start_rowx=0)  # <== lists of all values in column 3 of input .xlsx
nopingsdoc = docx.Document()
nopingsdocAbsPath = input("Please input absolute path to the file where unreachable nodes should be stored (like C:\\Users\\boburciu\\Downloads\\problematicNodes.docx ):")

vmatch = ["VM", "CloudLens", "centos", "CentOS", "ANVL", "VLM"]    # exclude nodes from ICMP test if their  names contain these values
reps_dict={}    # a dict of key=representative/owner and value=len(dict), so that each rep gets associated a sticky value

for i in range(len(names)):
        if any(x in str(names[i]) for x in vmatch):   # check if multiple strings exist in another string, use "all" instead of "any" if needed to verify all at once
                continue
        else:
                print(f"\n Checking host {ips[i]}")
                pkt=sr1(IP(dst=ips[i])/ICMP(id=1,seq=13983),timeout=2)
                if pkt == None:
                        if len(reps[i])==0 :
                                reps[i]="no_owner_registered"    # when there's no representative/owner (no value in input .xlsx in 3rd column), replace it

                        if reps[i] in reps_dict:
                                print(f"The host {ips[i]}:'{names[i]}' is down. Node owner '{reps[i]}', please help on-site resources with node's physical position.")
                                nopingsdoc.add_paragraph().add_run(f"The host {ips[i]}:'{names[i]}' is down. Node owner '{reps[i]}', please help on-site resources with node's physical position.").font.color.rgb = RGBColor((reps_dict[reps[i]]*100)%255, ((reps_dict[reps[i]]**3)*10)%255, ( 255-5*(reps_dict[reps[i]]**3) )%255 )  # write "some text" to file, colored by RGB(z%255, pow(z,3)%255, (255-5*pow(z,3))%155
                                nopingsdoc.save(nopingsdocAbsPath)
                        else:
                                reps_dict[reps[i]]=len(reps_dict)      # dict of key=rep/owner and value=0..6
                                print(f"The host {ips[i]}:'{names[i]}' is down. Node owner '{reps[i]}', please help on-site resources with node's physical position.")
                                nopingsdoc.add_paragraph().add_run(f"The host {ips[i]}:'{names[i]}' is down. Node owner '{reps[i]}', please help on-site resources with node's physical position.").font.color.rgb = RGBColor((reps_dict[reps[i]]*100)%255, ((reps_dict[reps[i]]**3)*10)%255, ( 255-5*(reps_dict[reps[i]]**3) )%255 )  # write "some text" to file, colored by RGB(z%255, pow(z,3)%255, (255-5*pow(z,3))%155)
                                nopingsdoc.save(nopingsdocAbsPath)
# -------------------------
# How to run:
# PS C:\Users\boburciu> python C:\Users\boburciu\Desktop\Automation_Python_for_Network_Engineers_by_David_Bombal\pingLabWin.py
# Please input absolute path to the file where unreachable nodes should be stored (like C:\Users\boburciu\Downloads\problematicNodes.docx ):C:\Users\boburciu\Downloads\problematicNodes.docx
# Checking host 10.38.161.2
# Begin emission:
# Finished sending 1 packets.
# :
# :
# -------------------------
# To use this script in Win PyCharm, you need xlrd scapy and docx modules installed, then Pycharm ALT+CRLS+S > Project Interpreter > + for modules
#
# PS C:\Users\boburciu> pip3 install xlrd
# Collecting xlrd
  # Downloading https://files.pythonhosted.org/packages/b0/16/63576a1a001752e34bf8ea62e367997530dc553b689356b9879339cf45a4
# /xlrd-1.2.0-py2.py3-none-any.whl (103kB)
     # |████████████████████████████████| 112kB 726kB/s
# Installing collected packages: xlrd
# Successfully installed xlrd-1.2.0
# WARNING: You are using pip version 19.2.3, however version 20.1.1 is available.
# You should consider upgrading via the 'python -m pip install --upgrade pip' command.
# PS C:\Users\boburciu>
# PS C:\Users\boburciu> pip3 install xlrd
# Requirement already satisfied: xlrd in c:\users\boburciu\appdata\local\programs\python\python37-32\lib\site-packages (1.
# 2.0)
# WARNING: You are using pip version 19.2.3, however version 20.1.1 is available.
# You should consider upgrading via the 'python -m pip install --upgrade pip' command.
# PS C:\Users\boburciu> pip3 install scapy
# Requirement already satisfied: scapy in c:\users\boburciu\appdata\local\programs\python\python37-32\lib\site-packages\scapy-git_archive.devef60120dfb-py3.7.egg (git-archive.devef60120dfb)
# WARNING: You are using pip version 19.2.3, however version 20.1.1 is available.
# You should consider upgrading via the 'python -m pip install --upgrade pip' command.
# PS C:\Users\boburciu>
# PS C:\Users\boburciu> pip3 install python-docx	# docx module is not for Python3
# Collecting python-docx
  # Downloading https://files.pythonhosted.org/packages/e4/83/c66a1934ed5ed8ab1dbb9931f1779079f8bca0f6bbc5793c06c4b5e7d671/python-docx-0.8.10.tar.gz (5.5MB)
     # |████████████████████████████████| 5.5MB 6.4MB/s
# Requirement already satisfied: lxml>=2.3.2 in c:\users\boburciu\appdata\local\programs\python\python37-32\lib\site-packages (from python-docx) (4.5.1)
# Installing collected packages: python-docx
  # Running setup.py install for python-docx ... done
# Successfully installed python-docx-0.8.10
# WARNING: You are using pip version 19.2.3, however version 20.1.1 is available.
# You should consider upgrading via the 'python -m pip install --upgrade pip' command.
# -------------------------

# >>> from scapy.all import *
# >>> p=sr(IP(dst="10.38.161.65")/ICMP(),timeout=2)
# Begin emission:
# ....Finished sending 1 packets.
# *
# Received 5 packets, got 1 answers, remaining 0 packets
# >>>
# Scapy generated frame:

# Frame 489: 42 bytes on wire (336 bits), 42 bytes captured (336 bits) on interface \Device\NPF_{62A356FA-03FF-4F6D-BDC3-0F778786F985}, id 0
# Ethernet II, Src: Cisco_3c:7a:00 (00:05:9a:3c:7a:00), Dst: CIMSYS_33:44:55 (00:11:22:33:44:55)
# Internet Protocol Version 4, Src: 10.22.219.99, Dst: 10.38.161.65
# Internet Control Message Protocol
    # Type: 8 (Echo (ping) request)
    # Code: 0
    # Checksum: 0xf7ff [correct]
    # [Checksum Status: Good]
    # Identifier (BE): 0 (0x0000)
    # Identifier (LE): 0 (0x0000)
    # Sequence number (BE): 0 (0x0000)
    # Sequence number (LE): 0 (0x0000)
    # [No response seen]
        # [Expert Info (Warning/Sequence): No response seen to ICMP request]
            # [No response seen to ICMP request]
            # [Severity level: Warning]
            # [Group: Sequence]

# sr1(IP(dst='10.38.161.65')/ICMP(id=1, seq=13983),timeout=2)
# Begin emission:
# Finished sending 1 packets.
# .*
# Received 2 packets, got 1 answers, remaining 0 packets
# <IP  version=4 ihl=5 tos=0x0 len=28 id=61843 flags= frag=0 ttl=60 proto=icmp chksum=0xfc6c src=10.38.161.65 dst=10.22.219.99 |<ICMP  type=echo-reply code=0 chksum=0xc95f id=0x1 seq=0x369f |>>

# sr1(IP(dst='10.38.161.65')/ICMP(id=1, seq=13983),timeout=2)
# Begin emission:
# Finished sending 1 packets.
# .*
# Received 2 packets, got 1 answers, remaining 0 packets
# <IP  version=4 ihl=5 tos=0x0 len=28 id=61843 flags= frag=0 ttl=60 proto=icmp chksum=0xfc6c src=10.38.161.65 dst=10.22.219.99 |<ICMP  type=echo-reply code=0 chksum=0xc95f id=0x1 seq=0x369f |>>

# Frame 4630: 42 bytes on wire (336 bits), 42 bytes captured (336 bits) on interface \Device\NPF_{62A356FA-03FF-4F6D-BDC3-0F778786F985}, id 0
# Ethernet II, Src: Cisco_3c:7a:00 (00:05:9a:3c:7a:00), Dst: CIMSYS_33:44:55 (00:11:22:33:44:55)
# Internet Protocol Version 4, Src: 10.22.219.99, Dst: 10.38.161.65
# Internet Control Message Protocol
    # Type: 8 (Echo (ping) request)
    # Code: 0
    # Checksum: 0xc15f [correct]
    # [Checksum Status: Good]
    # Identifier (BE): 1 (0x0001)
    # Identifier (LE): 256 (0x0100)
    # Sequence number (BE): 13983 (0x369f)
    # Sequence number (LE): 40758 (0x9f36)
    # [Response frame: 4631]
	
# -------------------------
# from docx import Document
# from docx.shared import RGBColor
#
# doc = Document()
# doc.add_paragraph().add_run('some text').font.color.rgb=RGBColor(0x99, 0x64, 0xa9)	# write "some text" to file, colored
# doc.save("C:\\Users\\boburciu\\Downloads\\problematicNodes.docx")
# doc.add_paragraph().add_run('2nd line of text').font.color.rgb=RGBColor(0x77, 0x64, 0xa9)
# doc.save("C:\\Users\\boburciu\\Downloads\\problematicNodes.docx")
